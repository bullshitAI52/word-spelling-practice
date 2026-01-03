# -*- coding: utf-8 -*-
"""
基于模板标记的 Word 表格批量提取工具 (tkinter) - 优化版
- 模板中用 {{标记名}} 标注要提取的单元格
- 解析模板所有表格，按 从上到下、从左到右 顺序记录 (表序, 行, 列, 标记名)
- 遍历 ./Files 下所有 .docx ，doc要转为docx, 在相同(表序, 行, 列) 处取值写入 ./汇总.xlsx
- UI 左：导入模板 / 提取文件标记信息 / 退出程序；右：Text 显示坐标与标记

优化点:
1.  重构代码结构，提升可读性和可维护性。
2.  提取过程采用多线程，防止 UI 卡顿。
3.  修复核心性能问题：避免在处理单个文件时重复解析同一个表格。
4.  使用 pathlib 进行路径管理，增强跨平台兼容性。
5.  实现跨平台的 "打开汇总表" 功能 (Windows, macOS, Linux)。
6.  引入更现代的 Python 写法和更清晰的异常处理。
"""
import os
import re
import sys
import traceback
import webbrowser
import threading
import queue
from pathlib import Path
from itertools import groupby
from tkinter import ttk, messagebox, filedialog, font as tkfont, Tk, Text, Frame, Scrollbar

# ----------------- 依赖检查 -----------------
try:
    from docx import Document
    from openpyxl import Workbook
except ImportError:
    missing = []
    try:
        import docx
    except ImportError:
        missing.append("python-docx")
    try:
        import openpyxl
    except ImportError:
        missing.append("openpyxl")
    if missing:
        raise SystemExit(f"缺少依赖：{', '.join(missing)}\n请先执行：pip install {' '.join(missing)}")

# ----------------- 常量定义 -----------------
BASE_DIR = Path(__file__).parent
FILES_DIR = BASE_DIR / "Files"
OUTPUT_XLSX = BASE_DIR / "汇总.xlsx"
MARK_PATTERN = re.compile(r"\{\{(.+?)\}\}")  # {{标记名}}
V_MERGE_CONTINUE = "continue"

# ----------------- Word 表解析核心逻辑 -----------------
class GridCell:
    """存储单元格的逻辑结构，处理合并单元格。"""
    __slots__ = ("anchor", "visible", "text", "rowspan", "colspan")
    def __init__(self, anchor=None, visible=False, text="", rowspan=1, colspan=1):
        self.anchor = anchor    # 合并区域的左上角坐标 (r, c)
        self.visible = visible  # 是否为合并区域的左上角单元格
        self.text = text
        self.rowspan = rowspan
        self.colspan = colspan

def _get_tcPr_prop(cell, name):
    """安全地获取单元格属性。"""
    try:
        return getattr(cell._tc.tcPr, name)
    except AttributeError:
        return None

def _to_int(val, default=1):
    """安全地将值转为整数。"""
    try:
        return int(str(val).strip())
    except (ValueError, TypeError):
        return default

def build_table_grid(table):
    """将 python-docx 的 table 对象解析为包含合并信息的逻辑网格。"""
    # 1. 确定网格的维度（考虑横向合并）
    col_counts = [sum(_to_int(getattr(_get_tcPr_prop(cell, "gridSpan"), "val", 1)) for cell in row.cells) for row in table.rows]
    n_rows = len(table.rows)
    n_cols = max(col_counts) if col_counts else 0
    if not n_rows or not n_cols:
        return [], 0, 0

    grid = [[None for _ in range(n_cols)] for _ in range(n_rows)]
    
    # 2. 遍历物理单元格，填充逻辑网格
    for r_idx, row in enumerate(table.rows):
        c_idx = 0
        for cell in row.cells:
            while c_idx < n_cols and grid[r_idx][c_idx] is not None:
                c_idx += 1
            if c_idx >= n_cols:
                break

            # 处理横向合并
            grid_span = _get_tcPr_prop(cell, "gridSpan")
            colspan = _to_int(getattr(grid_span, "val", 1)) if grid_span else 1

            # 处理纵向合并
            rowspan = 1
            v_merge = _get_tcPr_prop(cell, "vMerge")
            is_continue = False
            if v_merge is not None:
                val = getattr(v_merge, "val", None)
                if val is None or str(val).lower() == V_MERGE_CONTINUE:
                    is_continue = True

            if is_continue:
                # 向上查找锚点并扩展其 rowspan
                for r_scan in range(r_idx - 1, -1, -1):
                    if grid[r_scan][c_idx] and grid[r_scan][c_idx].anchor:
                        anchor_r, anchor_c = grid[r_scan][c_idx].anchor
                        if grid[anchor_r][anchor_c]:
                            grid[anchor_r][anchor_c].rowspan += 1
                        break
            
            # 填充被当前单元格（及其合并）所占据的网格区域
            text = cell.text.replace("\n", " ").strip()
            anchor_pos = (r_idx, c_idx)
            for i in range(rowspan):
                for j in range(colspan):
                    if r_idx + i < n_rows and c_idx + j < n_cols:
                        is_anchor_cell = (i == 0 and j == 0)
                        grid[r_idx + i][c_idx + j] = GridCell(
                            anchor=anchor_pos,
                            visible=is_anchor_cell,
                            text=text if is_anchor_cell else "",
                            rowspan=rowspan if is_anchor_cell else 1,
                            colspan=colspan if is_anchor_cell else 1
                        )
            c_idx += colspan

    # 3. 回填 rowspan (因为 continue 是向后看的)
    for r in range(n_rows - 2, -1, -1):
        for c in range(n_cols):
            cell = grid[r][c]
            if cell and cell.visible:
                # 检查正下方的单元格是否是合并的延续部分
                below_cell = grid[r + 1][c]
                if below_cell and not below_cell.visible and below_cell.anchor == cell.anchor:
                    # 找到合并链的底部，以确定总的 rowspan
                    final_rowspan = 1
                    for r_scan in range(r + 1, n_rows):
                        scan_cell = grid[r_scan][c]
                        if scan_cell and scan_cell.anchor == cell.anchor:
                            final_rowspan += 1
                        else:
                            break
                    cell.rowspan = final_rowspan

    # 4. 填补可能因表格结构错误留下的空位
    for r in range(n_rows):
        for c in range(n_cols):
            if grid[r][c] is None:
                grid[r][c] = GridCell()

    return grid, n_rows, n_cols

# ----------------- 业务逻辑 -----------------
class WordExtractor:
    @staticmethod
    def collect_marks_from_template(doc_path: Path):
        """从模板 Word 文档中收集所有标记及其坐标。"""
        if not doc_path.is_file():
            raise FileNotFoundError(doc_path)
        
        doc = Document(doc_path)
        marks = []
        for ti, table in enumerate(doc.tables):
            grid, n_rows, n_cols = build_table_grid(table)
            for r in range(n_rows):
                for c in range(n_cols):
                    gc = grid[r][c]
                    if gc and gc.visible and gc.text:
                        for m in MARK_PATTERN.finditer(gc.text):
                            name = m.group(1).strip()
                            if name:
                                marks.append({"table": ti, "row": r, "col": c, "name": name})
        return marks

    @staticmethod
    def extract_data(files: list[Path], marks: list[dict], status_callback):
        """从文件列表中根据标记提取数据。"""
        # 按表索引对标记进行分组，以优化性能
        marks.sort(key=lambda m: m['table'])
        grouped_marks = {ti: list(g) for ti, g in groupby(marks, key=lambda m: m['table'])}
        
        rows_out = []
        total_files = len(files)

        for i, path in enumerate(files):
            status_callback(f"正在处理: {i+1}/{total_files} - {path.name}")
            try:
                doc = Document(path)
                tables = doc.tables
            except Exception as e:
                print(f"无法读取文件 {path.name}: {e}")
                rows_out.append([path.stem] + [""] * len(marks))
                continue

            values = {}
            for ti, table_marks in grouped_marks.items():
                if ti >= len(tables):
                    continue
                
                grid, n_rows, n_cols = build_table_grid(tables[ti])
                if not grid:
                    continue

                for m in table_marks:
                    r, c = m["row"], m["col"]
                    val = ""
                    if 0 <= r < n_rows and 0 <= c < n_cols:
                        gc = grid[r][c]
                        # 如果坐标落在被覆盖区域，取其锚点文本
                        anchor_r, anchor_c = gc.anchor if gc.anchor else (r, c)
                        if 0 <= anchor_r < n_rows and 0 <= anchor_c < n_cols:
                            val = grid[anchor_r][anchor_c].text
                    
                    # 清理文本中可能残留的 {{...}}
                    values[m['name']] = MARK_PATTERN.sub("", val).strip()

            # 按原始标记顺序排列结果
            ordered_values = [values.get(m["name"], "") for m in marks]
            rows_out.append([path.stem] + ordered_values)

        # 过滤掉所有数据列都为空的行
        filtered_rows = [row for row in rows_out if any(cell for cell in row[1:])]
        return filtered_rows

    @staticmethod
    def save_to_excel(data: list[list], headers: list[str], output_path: Path):
        """将数据保存到 Excel 文件。"""
        wb = Workbook()
        ws = wb.active
        ws.title = "提取结果"
        ws.append(headers)
        for row in data:
            ws.append(row)
        wb.save(output_path)

# ----------------- UI 界面 -----------------
class App(Tk):
    def __init__(self):
        super().__init__()
        self.title("Word 表格批量提取 (优化版 by Gemini)")
        self.geometry("650x420")

        self._setup_styles()
        
        self.marks = []
        self.extractor = WordExtractor()
        self.queue = queue.Queue()

        self._build_ui()
        self._process_queue()

        FILES_DIR.mkdir(exist_ok=True)

    def _setup_styles(self):
        style = ttk.Style(self)
        try:
            font_family = "Microsoft YaHei"
            default_font = tkfont.nametofont("TkDefaultFont")
            text_font = tkfont.nametofont("TkTextFont")
            fixed_font = tkfont.nametofont("TkFixedFont")
            for f in (default_font, text_font, fixed_font):
                f.configure(family=font_family, size=11)
        except Exception:
            pass  # Font setup is not critical
        style.configure("TButton", padding=(10, 6))
        style.configure("TLabel", padding=(2, 2))
        style.configure("TFrame", padding=(8, 8))

    def _build_ui(self):
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, minsize=200, weight=0)
        self.grid_columnconfigure(1, weight=1)

        # --- 左栏 ---
        left = ttk.Frame(self)
        left.grid(row=0, column=0, sticky="nswe")
        
        ttk.Button(left, text="1. 导入模板 (.docx)", command=self.on_load_template).pack(fill="x", pady=5)
        ttk.Button(left, text="2. 提取数据 (从 Files 目录)", command=self.on_extract_all).pack(fill="x", pady=5)
        ttk.Button(left, text="打开汇总表", command=self.open_xlsx).pack(fill="x", pady=5, side="bottom")
        
        tip = ("使用说明：\n"
               "1. 点击“导入模板”，选择含 {{标记}} 的 DOCX。\n"
               "2. 右侧将列出 表序/坐标/标记。\n"
               "3. 将要处理的 .docx 文件放入 Files 目录。\n"
               "4. 点击“提取数据”，程序将自动处理并生成\n   “汇总.xlsx”。")
        ttk.Label(left, text=tip, justify="left", wraplength=180).pack(pady=20, fill="x")

        # --- 右侧 ---
        right = ttk.Frame(self)
        right.grid(row=0, column=1, sticky="nsew", padx=(5, 0))
        right.grid_rowconfigure(1, weight=1)
        right.grid_columnconfigure(0, weight=1)

        ttk.Label(right, text="模板标记预览 (从上到下, 从左到右)").grid(row=0, column=0, sticky="w", pady=(0, 4))
        self.text = Text(right, wrap="word", height=10)
        self.text.grid(row=1, column=0, sticky="nsew")
        ybar = ttk.Scrollbar(right, orient="vertical", command=self.text.yview)
        self.text.configure(yscrollcommand=ybar.set)
        ybar.grid(row=1, column=1, sticky="ns")

        # --- 底部状态栏 ---
        self.status = ttk.Label(self, text="准备就绪", anchor="w")
        self.status.grid(row=1, column=0, columnspan=2, sticky="ew", padx=5, pady=(2, 2))

    def set_status(self, msg):
        self.status.config(text=msg)

    def _process_queue(self):
        """处理来自工作线程的消息队列。"""
        try:
            msg, data = self.queue.get_nowait()
            if msg == "status":
                self.set_status(data)
            elif msg == "done":
                self.set_status(f"提取完成！已写入：{OUTPUT_XLSX.name}")
                messagebox.showinfo("完成", f"提取完成，共处理 {data} 个文件。\n已写入：\n{OUTPUT_XLSX}")
            elif msg == "error":
                self.set_status("出现错误")
                messagebox.showerror("错误", f"处理失败：\n{data}")
        except queue.Empty:
            pass
        finally:
            self.after(100, self._process_queue)

    def on_load_template(self):
        path_str = filedialog.askopenfilename(
            title="选择模板（包含 {{标记}} 的 DOCX）",
            filetypes=[("Word 文档", "*.docx")]
        )
        if not path_str:
            return

        template_path = Path(path_str)
        try:
            self.set_status("正在解析模板...")
            self.marks = self.extractor.collect_marks_from_template(template_path)
            self.show_marks()
            if not self.marks:
                messagebox.showwarning("提示", "未在模板的表格里找到 {{标记}}。")
                self.set_status("模板中未找到标记")
            else:
                messagebox.showinfo("完成", f"模板解析完成，共找到 {len(self.marks)} 个标记。")
                self.set_status(f"模板加载成功: {template_path.name}")
        except Exception as e:
            self.set_status("解析失败")
            messagebox.showerror("错误", f"解析模板失败：\n{e}")
            traceback.print_exc()

    def show_marks(self):
        self.text.delete("1.0", "end")
        if not self.marks:
            self.text.insert("end", "尚未加载模板或未识别到标记。\n")
            return
        lines = [f"[{i:02d}] T{m['table']+1} R{m['row']+1}C{m['col']+1} -> {{ {m['name']} }}" for i, m in enumerate(self.marks, 1)]
        self.text.insert("end", "\n".join(lines))

    def on_extract_all(self):
        if not self.marks:
            messagebox.showinfo("提示", "请先导入包含 {{标记}} 的模板。")
            return
        
        docx_paths = list(FILES_DIR.glob("*.docx"))
        if not docx_paths:
            messagebox.showinfo("提示", f"{FILES_DIR.name} 目录中没有 .docx 文件。")
            return

        self.set_status("开始提取数据...")
        # 在新线程中运行提取任务
        thread = threading.Thread(target=self._run_extraction, args=(docx_paths, self.marks), daemon=True)
        thread.start()

    def _run_extraction(self, docx_paths, marks):
        """在工作线程中执行的提取和保存逻辑。"""
        try:
            def status_callback(msg):
                self.queue.put(("status", msg))

            # 1. 提取数据
            extracted_data = self.extractor.extract_data(docx_paths, marks, status_callback)
            
            # 2. 准备保存
            headers = ["文件名"] + [m["name"] for m in marks]
            
            # 3. 保存到 Excel
            self.queue.put(("status", f"正在保存到 {OUTPUT_XLSX.name}..."))
            self.extractor.save_to_excel(extracted_data, headers, OUTPUT_XLSX)
            
            self.queue.put(("done", len(docx_paths)))
        except Exception as e:
            traceback.print_exc()
            self.queue.put(("error", str(e)))

    def open_xlsx(self):
        if not OUTPUT_XLSX.is_file():
            messagebox.showwarning("文件不存在", f"汇总表尚未生成：\n{OUTPUT_XLSX}")
            return
        try:
            # 使用 webbrowser 以实现跨平台打开文件
            webbrowser.open(OUTPUT_XLSX.as_uri())
        except Exception as e:
            messagebox.showerror("打开失败", f"无法打开文件：\n{e}")

if __name__ == "__main__":
    try:
        app = App()
        app.mainloop()
    except SystemExit as e:
        print(e, file=sys.stderr)
    except Exception:
        traceback.print_exc()
