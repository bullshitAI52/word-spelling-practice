import json
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
import openpyxl


# --- WordTableParser Class ---
class WordTableParser:
    def __init__(self, filepath: str):
        self.filepath = filepath
        self.document = None
        self.tables = []

    def load_document(self) -> None:
        """加载 Word 文件"""
        self.document = Document(self.filepath)

    def extract_tables(self):
        """提取所有表格，转换为二维数组"""
        if self.document is None:
            self.load_document()

        self.tables = []
        for table in self.document.tables:
            parsed_table = []
            for row in table.rows:
                parsed_row = [cell.text.strip() for cell in row.cells]
                parsed_table.append(parsed_row)
            self.tables.append(parsed_table)
        return self.tables

    def save_as_json(self, output_path: str) -> None:
        if not self.tables:
            self.extract_tables()
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(self.tables, f, ensure_ascii=False, indent=2)

    def save_as_excel(self, output_path: str) -> None:
        if not self.tables:
            self.extract_tables()
        wb = openpyxl.Workbook()
        wb.remove(wb.active)
        for idx, table in enumerate(self.tables, start=1):
            ws = wb.create_sheet(f"Table{idx}")
            for row in table:
                ws.append(row)
        wb.save(output_path)

    def save_as_html(self, output_path: str) -> None:
        if not self.tables:
            self.extract_tables()

        html_content = "<html><head><meta charset='utf-8'></head><body>\n"
        for idx, table in enumerate(self.tables, start=1):
            html_content += f"<h3>Table {idx}</h3>\n<table border='1' cellspacing='0' cellpadding='5'>\n"
            for row in table:
                html_content += "<tr>" + "".join(f"<td>{cell}</td>" for cell in row) + "</tr>\n"
            html_content += "</table><br>\n"
        html_content += "</body></html>"

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)


# --- WordParserUI Class ---
class WordParserUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Word 表格转换工具 (JSON / Excel / HTML)")
        self.root.geometry("600x250")

        self.input_file = None
        self.output_file = None
        self.output_format = tk.StringVar(value="JSON")
        
        main_frame = ttk.Frame(root, padding="15 15 15 15")
        main_frame.pack(fill='both', expand=True)
        main_frame.columnconfigure(1, weight=1)

        # 1. 输入文件
        ttk.Label(main_frame, text="Word 文件:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.btn_open = ttk.Button(main_frame, text="选择...", command=self.select_input, width=10)
        self.btn_open.grid(row=0, column=2, padx=5, pady=5)
        self.input_path_var = tk.StringVar()
        self.input_entry = ttk.Entry(main_frame, textvariable=self.input_path_var, state='readonly')
        self.input_entry.grid(row=0, column=1, padx=5, pady=5, sticky="ew")

        # 2. 输出格式
        ttk.Label(main_frame, text="输出格式:").grid(row=1, column=0, padx=5, pady=10, sticky="w")
        self.format_menu = ttk.Combobox(main_frame, textvariable=self.output_format, 
                                        values=["JSON", "Excel", "HTML"], state="readonly", width=10)
        self.format_menu.grid(row=1, column=1, padx=5, pady=10, sticky="w")
        self.format_menu.bind("<<ComboboxSelected>>", self.clear_output_path)

        # 3. 输出文件
        ttk.Label(main_frame, text="保存路径:").grid(row=2, column=0, padx=5, pady=5, sticky="w")
        self.btn_save = ttk.Button(main_frame, text="选择...", command=self.select_output, width=10)
        self.btn_save.grid(row=2, column=2, padx=5, pady=5)
        self.output_path_var = tk.StringVar()
        self.output_entry = ttk.Entry(main_frame, textvariable=self.output_path_var, state='readonly')
        self.output_entry.grid(row=2, column=1, padx=5, pady=5, sticky="ew")

        # 4. 执行按钮
        self.btn_extract = ttk.Button(main_frame, text="开始转换", command=self.extract, style='Accent.TButton')
        self.btn_extract.grid(row=3, column=0, columnspan=3, pady=20, sticky="s")
        
        style = ttk.Style()
        style.theme_use('clam')

    def clear_output_path(self, event=None):
        self.output_file = None
        self.output_path_var.set("")

    def select_input(self):
        file_path = filedialog.askopenfilename(
            title="选择 Word 文件", 
            filetypes=[("Word 文件", "*.docx")]
        )
        if file_path:
            self.input_file = file_path
            self.input_path_var.set(file_path)

    def select_output(self):
        ext_map = {"JSON": ".json", "Excel": ".xlsx", "HTML": ".html"}
        fmt = self.output_format.get()
        file_path = filedialog.asksaveasfilename(
            title="保存文件",
            defaultextension=ext_map[fmt],
            filetypes=[(f"{fmt} 文件", f"*{ext_map[fmt]}")]
        )
        if file_path:
            self.output_file = file_path
            self.output_path_var.set(file_path)

    def extract(self):
        if not self.input_file:
            messagebox.showerror("错误", "请先选择输入的 Word 文件！")
            return
        if not self.output_file:
            messagebox.showerror("错误", "请先选择输出文件路径！")
            return
        
        self.btn_extract.config(state=tk.DISABLED, text="正在转换...")

        try:
            parser = WordTableParser(self.input_file)
            parser.extract_tables()

            fmt = self.output_format.get()
            if fmt == "JSON":
                parser.save_as_json(self.output_file)
            elif fmt == "Excel":
                parser.save_as_excel(self.output_file)
            elif fmt == "HTML":
                parser.save_as_html(self.output_file)

            messagebox.showinfo("成功", f"表格已转换并保存为 {fmt} 文件：\n{self.output_file}")
        except Exception as e:
            messagebox.showerror("失败", f"转换失败：\n请确保文件格式正确且未被其他程序占用。\n详细错误：{e}")
        finally:
            self.btn_extract.config(state=tk.NORMAL, text="开始转换")


if __name__ == "__main__":
    root = tk.Tk()
    app = WordParserUI(root)
    root.mainloop()